"""
Order Processing System

This script automates order processing for an e-commerce business. It provides functionality for:
- Loading and processing order data from Excel files
- Generating shipping labels, invoices, and order documents
- Handling different payment and delivery methods
- Processing product images
- Creating PDF and DOCX format documents

The system reads orders from an Excel file, processes each order based on its status
and delivery method, and generates appropriate documentation.

Author: Rytis Bimbiras
Last Updated: May 3, 2025
"""

# ---------------------------------------------------------------------------
# QUICK START OVERVIEW
#
# This script is designed to be run as a stand‑alone utility. When executed,
# `process_all_orders()` is called from the `if __name__ == "__main__":`
# guard at the very bottom.  The high‑level flow is:
#
#     Excel → pandas.DataFrame  → per‑order grouping → Word/PDF generation
#                          ↑                         ↓
#                      status flag            updated Excel saved
#
# Each helper in the “Utility”, “Image Processing”, “PDF Generation”, and
# “Document Creation” sections manipulates a **single, well‑defined piece**
# of that pipeline.  The comments sprinkled throughout the functions below
# try to explain *why* something is happening – not just *what*.
# ---------------------------------------------------------------------------

# Import necessary libraries for file handling, data processing, and document generation
import os                               # For file and directory operations
import pandas as pd                     # For Excel file manipulation and data processing
import requests                         # For making HTTP requests to download images
from io import BytesIO                  # For handling binary data streams (images)
from PIL import Image, ImageOps, UnidentifiedImageError # For image processing operations
from datetime import datetime # For date and time operations

# Import libraries for PDF generation
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage
from reportlab.lib import colors # For color definitions in PDFs
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle # For text styling in PDFs
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT # For text alignment options
from reportlab.lib.pagesizes import A4, A7 # For standard page size definitions
from reportlab.pdfbase.ttfonts import TTFont # For custom font handling
from reportlab.pdfbase import pdfmetrics # For font registration
from reportlab.pdfgen import canvas # For low-level PDF operations

# Import libraries for Word document generation
from docx import Document # For creating Word documents
from docx.shared import Inches # For measurement units in Word docs

# Import libraries for barcode generation
from reportlab.graphics.barcode import eanbc # For EAN barcode generation
from reportlab.graphics.shapes import Drawing # For vector drawing operations

# Register fonts with proper bold support for PDF documents
pdfmetrics.registerFont(TTFont('DejaVuSans', 'DejaVuSans.ttf'))
pdfmetrics.registerFont(TTFont('DejaVuSans-Bold', 'DejaVuSans-Bold.ttf'))

# =========================
# Configuration and Constants
# =========================

# Dictionary containing information about different shop locations
# Each shop has an email address and physical address
SHOP_INFO = {
    "k_mega": {
        "email": "mega@winnersport.lt",
        "address": "Islandijos pl. 32, Kaunas, LT-47446"
    },
    "v_outlet": {
        "email": "outlet@winnersport.lt",
        "address": "Vytauto Pociūno g. 8, Vilnius, LT-06231"
    },
    "v_ozas": {
        "email": "ozas@winnersport.lt",
        "address": "Ozo g. 18, Vilnius, LT-08243"
    }
}

# Company information used for invoices and shipping documents
COMPANY_INFO = {
    "name": 'JSC "Briedziukas"',
    "phone": "+37061402705",
    "email": "info@winnersport.lt",
    "account_no": "LT343500010001083590",
    "bank": "Paysera LT",
    "bank_code": "35000",
    "swift_code": "EVIULT2VXXX",
    "company_code": "122571351",
    "vat_code": "LT225713515"
}

# =========================
# Utility Functions
# =========================

#
# ------------------------------------------------------------
# Format phone numbers into a consistent international style
# ------------------------------------------------------------
def standardize_phone_number(phone):
    """
    Standardizes phone number format to international format with country code.

    Args:
        phone (str or int): The phone number to format
    Returns:
        str: Formatted phone number starting with '+' and country code
    """
    # Normalise the incoming phone value – users sometimes paste it with
    # spaces or even numeric types; we turn it into a clean string first.
    # Convert phone to string and remove any spaces
    phone_str = str(phone).strip().replace(" ", "")
    
    # If phone starts with country code 370 (Lithuania), add + prefix
    if phone_str.startswith("370"):
        return f"+{phone_str}"
    
    # If phone is 8 digits (local Lithuanian number), add country code
    if len(phone_str) == 8 and not phone_str.startswith("+"):
        return f"+370{phone_str}"
    
    # Return original phone if it doesn't match specific patterns
    return phone_str

#
# ------------------------------------------------------------
# Fetch email or physical address for a given shop identifier
# ------------------------------------------------------------
def get_shop_info(shop, info_type=None):
    """
    Retrieves information about a specific shop.

    Args:
        shop (str): Shop identifier (e.g., 'k_mega', 'v_outlet')
        info_type (str, optional): Type of info to retrieve ('email' or 'address')
    Returns:
        dict or str: Shop information dictionary or specific info type value
    """
    # Get shop information from the SHOP_INFO constant
    shop_info = SHOP_INFO.get(shop, {})
    
    # Return specific information type if requested
    if info_type == 'email':
        return shop_info.get("email", "unknown@winnersport.lt") # Default email if not found
    elif info_type == 'address':
        return shop_info.get("address", "Unknown address") # Default address if not found
    
    # Return all shop info if no specific type requested
    return shop_info

# =========================
# Data Loading & Checking
# =========================

#
# ------------------------------------------------------------
# Read the raw Excel spreadsheet and return a cleaned DataFrame
# ------------------------------------------------------------
def load_orders_from_excel(filename):
    """
    Loads and prepares order data from an Excel file.

    Args:
        filename (str): Path to the Excel file
    Returns:
        pandas.DataFrame: DataFrame containing order data
    """
    # Read the Excel file into a pandas DataFrame
    orders_dataframe = pd.read_excel(filename)
    
    # Strip whitespace from column names for consistency
    orders_dataframe.columns = orders_dataframe.columns.str.strip()
    
    return orders_dataframe

#
# ------------------------------------------------------------
# Quick check to see if the given order is already marked done
# ------------------------------------------------------------
def is_order_already_processed(order_items):
    """
    Checks if an order has already been marked as processed.

    Args:
        order_items (pandas.DataFrame): DataFrame containing order items
    Returns:
        bool: True if order is already processed, False otherwise
    """
    # Get the 'order_processed' value from the first row, convert to string, strip whitespace,
    # convert to lowercase, and check if it equals "yes"
    return str(order_items.iloc[0].get("order_processed", "")).strip().lower() == "yes"

#
# ------------------------------------------------------------
# Flag an order as processed so we don't duplicate work later
# ------------------------------------------------------------
def mark_order_as_processed(orders_dataframe, order_id):
    """
    Marks an order as processed in the DataFrame.

    Args:
        orders_dataframe (pandas.DataFrame): DataFrame containing orders
        order_id: ID of the order to mark as processed
    """
    # Update 'order_processed' column to 'yes' for rows with matching order_id
    orders_dataframe.loc[orders_dataframe["order_id"] == order_id, "order_processed"] = "yes"

# =========================
# Image Processing
# =========================

#
# ------------------------------------------------------------
# Download product photo, resize, pad, border, return PIL image
# ------------------------------------------------------------
def download_and_process_image(image_url, size=200, border_width=5, border_color="black"):
    """
    Downloads and processes a product image from a URL.

    Args:
        image_url (str): URL of the image to download
        size (int): Target size (width/height) for the image
        border_width (int): Width of border to add around image
        border_color (str): Color of the border
    Returns:
        PIL.Image.Image or None: Processed image object or None if processing failed
    """
    # Check if image URL is valid
    if not pd.notna(image_url) or not isinstance(image_url, str) or not image_url.strip():
        return None # Return None if URL is invalid
    
    try:
        # Download the product photo (fail fast after 10 s to avoid long stalls)
        response = requests.get(image_url, timeout=10)
        response.raise_for_status() # Raise exception for unsuccessful HTTP requests
        
        # Open image from response content and convert to RGB format
        image = Image.open(BytesIO(response.content)).convert("RGB")
        
        # Preserve aspect ratio while shrinking so neither side exceeds *size* px
        image.thumbnail((size, size))
        
        # Create new white background image of target size
        image_with_padding = Image.new("RGB", (size, size), color="white")
        
        # Paste resized image onto white background, centered
        image_with_padding.paste(image, ((size - image.width) // 2, (size - image.height) // 2))
        
        # Add border around the image
        image_with_frame = ImageOps.expand(image_with_padding, border=border_width, fill=border_color)
        
        return image_with_frame
    
    except (requests.exceptions.RequestException, UnidentifiedImageError) as exception:
        # Handle network errors and image processing errors
        print(f"Error processing image: {exception}")
        return None

#
# ------------------------------------------------------------
# Safely add an image (or placeholder text) into the Word doc
# ------------------------------------------------------------
def attach_image_to_document(document, image, order_id, description=None):
    """
    Attaches an image to a Word document, with error handling.

    Args:
        document (docx.Document): Word document to attach image to
        image (PIL.Image.Image): Image to attach
        order_id: ID of the order (used for temporary file naming)
        description (str, optional): Description text if image cannot be attached
    """
    # Check if image is available
    if image is None:
        # Add warning message with description if image is not available
        if description:
            document.add_paragraph(f"⚠️ {description}")
        return
    
    # Create a temporary file path for the image
    image_path = f"temp_{order_id}_processed.jpg"
    
    # Save image to temporary file
    image.save(image_path)
    
    # Add image to document with width of 2 inches
    document.add_picture(image_path, width=Inches(2))
    
    # Remove temporary image file
    os.remove(image_path)

# =========================
# PDF Generation
# =========================

#
# ------------------------------------------------------------
# Build a compact A7 PDF shipping label with barcode & COD info
# ------------------------------------------------------------
def generate_shipping_label_pdf(filename, order_item, payment_status):
    """
    Creates a shipping label PDF for an order item.

    Args:
        filename (str): Output PDF filename
        order_item (dict): Dictionary containing order item details
        payment_status (str): Payment status (affects COD amount display)
    """
    # ------- Recipient / Sender information block ---------------------
    # Get shop identifier from order item
    shop = order_item.get("shop", "")
    
    # Get shop's address using shop identifier
    shop_address = get_shop_info(shop, 'address')
    
    # Define text styles for the document
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="LeftBold", alignment=TA_LEFT, fontName="DejaVuSans-Bold", fontSize=9, leading=11))
    styles.add(ParagraphStyle(name="Left", alignment=TA_LEFT, fontName="DejaVuSans", fontSize=9, leading=11))
    
    # Create PDF document with A7 size (small for shipping label)
    document = SimpleDocTemplate(filename, pagesize=A7, leftMargin=10, rightMargin=10, topMargin=10, bottomMargin=10)
    
    # List to hold all elements in the document
    elements = []
    
    # Add recipient information to shipping label
    elements.append(Paragraph(f"Ship to:", styles["LeftBold"]))
    elements.append(Paragraph(order_item.get('client', ''), styles["Left"]))
    elements.append(Paragraph(order_item.get("address", ""), styles["Left"]))
    elements.append(Paragraph(f"Tel: {standardize_phone_number(order_item.get('phone_number', ''))}", styles["Left"]))
    elements.append(Paragraph(order_item.get("email", ""), styles["Left"]))
    elements.append(Spacer(1, 12)) # Add vertical space
    
    # Add sender information to shipping label
    elements.append(Paragraph("From:", styles["LeftBold"]))
    elements.append(Paragraph(COMPANY_INFO["name"], styles["LeftBold"]))
    elements.append(Paragraph(shop_address, styles["Left"]))
    elements.append(Paragraph(f"Tel: {COMPANY_INFO['phone']}", styles["Left"]))
    elements.append(Paragraph(COMPANY_INFO["email"], styles["Left"]))
    elements.append(Spacer(1, 12)) # Add vertical space
    
    # If payment hasn’t been received, print COD so the courier knows the
    # amount to collect.
    # Add COD (Cash On Delivery) amount if payment wasn't received yet
    if payment_status != "received":
        try:
            # Calculate COD amount (price × quantity)
            price = float(order_item.get("item_price", 0))
            qty = float(order_item.get("ordered_quantity", 1))
            cod_text = f"COD Amount: {price * qty:.2f} €"
            elements.append(Paragraph(cod_text, styles["LeftBold"]))
            elements.append(Spacer(1, 12)) # Add vertical space
        except Exception:
            pass # Skip COD amount if calculation fails
    
    # ------- 1‑D Barcode (EAN‑13) creation -----------------------------
    # Get barcode value from order item
    barcode_input = order_item.get("barcode", "")
    
    try:
        # Format barcode number as required for EAN13 (13 digits)
        barcode_str = f"{float(barcode_input):.0f}".zfill(13)
        
        # Validate barcode length
        if len(barcode_str) != 13:
            raise ValueError("Invalid barcode length")
        
        # Create EAN13 barcode
        barcode = eanbc.Ean13BarcodeWidget(barcode_str)
        barcode.barWidth = 1.5 # Set bar width
        barcode.barHeight = 30 # Set bar height
        
        # Create drawing to hold barcode
        barcode_drawing = Drawing(100, 40)
        barcode_drawing.add(barcode)
        barcode_drawing.hAlign = 'LEFT'
        
        # Add space before barcode
        elements.append(Spacer(1, 8))
        
        # Add barcode to document
        elements.append(barcode_drawing)
    
    except Exception:
        # Add error message if barcode generation fails
        elements.append(Paragraph("⚠️ Invalid Barcode", styles["Left"]))
    
    # Build final PDF with all elements
    document.build(elements)

#
# ------------------------------------------------------------
# Assemble the logo + seller/buyer blocks for invoice PDFs
# ------------------------------------------------------------
def create_invoice_header(first_order_item, styles, doc_type="Invoice"):
    """
    Creates the header section for an invoice or similar document.

    Args:
        first_order_item (dict): First item in the order (contains order details)
        styles (dict): Dictionary of paragraph styles
        doc_type (str): Document type ("Invoice" or "Proforma Invoice")
    Returns:
        list: List of document elements for the header
    """
    # Initialize list for document elements
    elements = []
    
    try:
        # Add company logo
        logo_url = "https://www.winnersport.lt/static/v2/img/logo.png"
        logo_resp = requests.get(logo_url, timeout=10)
        logo_image_bytes = BytesIO(logo_resp.content)
        logo_pil_image = Image.open(logo_image_bytes)
        
        # Calculate logo size with proper aspect ratio
        max_width = 220
        max_height = 60
        ratio = min(max_width / logo_pil_image.width, max_height / logo_pil_image.height)
        logo_width = logo_pil_image.width * ratio
        logo_height = logo_pil_image.height * ratio
        
        # Reset BytesIO position to beginning
        logo_image_bytes.seek(0)
        
        # Create ReportLab image from logo bytes
        logo_image = RLImage(logo_image_bytes, width=logo_width, height=logo_height)
        logo_image.hAlign = 'LEFT'
        
        # Add logo to document
        elements.append(logo_image)
    
    except Exception:
        # Skip logo if there's any error
        pass
    
    # Add vertical space after logo
    elements.append(Spacer(1, 14))
    
    # Extract client information from order item
    client = first_order_item["client"]
    phone = standardize_phone_number(first_order_item["phone_number"])
    email = first_order_item.get("email", "info@client.lt")
    address = first_order_item["address"]
    
    try:
        # Format order date as YYYY-MM-DD
        order_date = pd.to_datetime(first_order_item["date"]).strftime("%Y-%m-%d")
    except Exception:
        # Use today's date if order date parsing fails
        order_date = datetime.today().strftime("%Y-%m-%d")
    
    # Calculate content width based on page size
    page_width = A4[0]
    content_width = page_width - 60
    
    # Create company information text block
    company_info = "<br/>".join([
        f'Seller: {COMPANY_INFO["name"]}',
        f'Account No.: {COMPANY_INFO["account_no"]}',
        f'Bank: {COMPANY_INFO["bank"]}',
        f'Bank code: {COMPANY_INFO["bank_code"]}',
        f'Swift code: {COMPANY_INFO["swift_code"]}',
        f'Company code: {COMPANY_INFO["company_code"]}',
        f'VAT code: {COMPANY_INFO["vat_code"]}',
        f'Tel.: {COMPANY_INFO["phone"]}',
        f'Email: {COMPANY_INFO["email"]}'
    ])
    
    # Create buyer information text block
    buyer_info = "<br/>".join([
        f'Buyer: {client}',
        f'Address: {address}',
        f'Tel.: {phone}',
        f'Email: {email}'
    ])
    
    # Create table with company and buyer information in two columns
    info_table = [[
        Paragraph(company_info, styles["Left"]),
        Paragraph(buyer_info, styles["Left"])
    ]]
    
    # Create and style the information table
    info = Table(info_table, colWidths=[content_width/2, content_width/2], hAlign='LEFT')
    info.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"), # Align content to top
        ("ALIGN", (0, 0), (-1, -1), "LEFT"), # Align content to left
        ("LEFTPADDING", (0, 0), (-1, -1), 0), # No left padding
        ("RIGHTPADDING", (0, 0), (-1, -1), 12), # Right padding of 12 points
    ]))
    
    # Add info table to document
    elements.append(info)
    
    # Add space after info table
    elements.append(Spacer(1, 12))
    
    # Get order ID and convert to integer
    order_id = first_order_item.get("order_id", "")
    
    # Add document title (Invoice or Proforma Invoice) with order number
    elements.append(Paragraph(f"{doc_type} nr. {int(float(order_id))}", styles["CenterBold"]))
    
    # Add issue date
    elements.append(Paragraph(f"Issue date: {order_date}", styles["Left"]))
    
    # Add space after header information
    elements.append(Spacer(1, 8))
    
    return elements

#
# ------------------------------------------------------------
# Turn order line items into a styled table and compute totals
# ------------------------------------------------------------
def create_invoice_items_table(order_items, styles):
    """
    Creates a table showing order items with prices and totals.

    Args:
        order_items (pandas.DataFrame): DataFrame containing order items
        styles (dict): Dictionary of paragraph styles
    Returns:
        tuple: (table, subtotal, vat, grand_total) - the table element and calculated totals
    """
    # Helper function to wrap text in a paragraph with appropriate style
    def wrap(text, bold=False):
        style = styles["LeftBold"] if bold else styles["Left"]
        return Paragraph(str(text), style)
    
    # Build the header first – every subsequent row will mimic this six‑column layout
    # Create table header row with column titles
    table_data = [[
        wrap("No.", True),
        wrap("Product", True),
        wrap("Qty", True),
        wrap("Price excl.VAT", True),
        wrap("Total excl.VAT", True),
        wrap("Total incl.VAT", True)
    ]]
    
    # Track running total of order value
    subtotal = 0
    
    # ------------------------------------------------------------------
    # Iterate over each order line to calculate per‑item pricing details
    # ------------------------------------------------------------------
    # Add a row for each order item
    for item_index, row in enumerate(order_items.itertuples(), 1):
        # Extract and calculate values
        qty = getattr(row, "ordered_quantity", 0) or 0
        original_price = getattr(row, "item_price", 0) or 0
        
        # Convert the store’s retail (incl. VAT) price to *net* price.
        # Lithuania’s VAT is 21 %, so dividing by 1.21 reverses it.
        # Calculate price excluding VAT (divide by 1.21 for 21% VAT)
        price_excl_vat = original_price / 1.21
        
        # Calculate row totals
        total_no_vat = qty * price_excl_vat
        vat_amount = total_no_vat * 0.21
        total_with_vat = total_no_vat + vat_amount
        
        # Add to order subtotal
        subtotal += total_no_vat
        
        # Extract product details
        item_name = getattr(row, "item_name", "")
        size = getattr(row, "size", "")
        barcode = getattr(row, "barcode", "")
        
        try:
            # Format barcode as integer if possible
            barcode = str(int(float(barcode))) if pd.notna(barcode) else ""
        except Exception:
            # Keep as string if conversion fails
            barcode = str(barcode)
        
        # Human‑friendly product description shown on the PDF row
        # Create product description with size and barcode if available
        product_str = f"{item_name}, size {size}, ({barcode})" if size else f"{item_name}, ({barcode})"
        
        # Add row to table
        table_data.append([
            wrap(str(item_index)), # Item number
            wrap(product_str), # Product description
            wrap(f"{qty:.0f}"), # Quantity
            wrap(f"{price_excl_vat:.2f} €"), # Price excluding VAT
            wrap(f"{total_no_vat:.2f} €"), # Total excluding VAT
            wrap(f"{total_with_vat:.2f} €") # Total including VAT
        ])
    
    # Calculate page and column widths
    page_width = A4[0]
    content_width = page_width - 60
    col_widths = [40, content_width-305, 40, 75, 75, 75]
    
    # Create table with calculated widths
    table = Table(table_data, colWidths=col_widths, hAlign='LEFT')
    
    # Apply styles to table
    table.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, 0), "DejaVuSans-Bold"), # Bold font for header row
        ("FONTNAME", (0, 1), (-1, -1), "DejaVuSans"), # Regular font for data rows
        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey), # Gray background for header
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey), # Grid lines
        ("ALIGN", (2, 0), (-1, 0), "CENTER"), # Center align header text
        ("ALIGN", (2, 1), (-1, -1), "RIGHT"), # Right align numeric values
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"), # Vertically center all content
        ("BOTTOMPADDING", (0, 0), (-1, 0), 6), # Bottom padding for header
        ("TOPPADDING", (0, 0), (-1, 0), 6), # Top padding for header
    ]))
    
    # Calculate VAT and grand total
    vat = subtotal * 0.21
    grand_total = subtotal + vat
    
    return table, subtotal, vat, grand_total

#
# ------------------------------------------------------------
# Bring header, item table, and totals together into a PDF file
# ------------------------------------------------------------
def generate_invoice_pdf(order_id, order_items, output_path, doc_type="Invoice"):
    """
    Generates a complete invoice PDF document.

    Args:
        order_id: ID of the order
        order_items (pandas.DataFrame): DataFrame containing order items
        output_path (str): Path to save the PDF file
        doc_type (str): Document type ("Invoice" or "Proforma Invoice")
    Returns:
        str: Path to the generated PDF file
    """
    # Define paragraph styles for the document
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="Left", alignment=TA_LEFT, fontName="DejaVuSans", fontSize=10, leading=13))
    styles.add(ParagraphStyle(name="LeftBold", alignment=TA_LEFT, fontName="DejaVuSans-Bold", fontSize=10))
    styles.add(ParagraphStyle(name="CenterBold", alignment=TA_CENTER, fontName="DejaVuSans-Bold", fontSize=14, spaceAfter=10))
    styles.add(ParagraphStyle(name="RightBold", alignment=TA_RIGHT, fontName="DejaVuSans-Bold", fontSize=10))
    styles.add(ParagraphStyle(name="Right", alignment=TA_RIGHT, fontName="DejaVuSans", fontSize=10))
    
    # Initialize list for document elements
    elements = []
    
    # Create and add document header
    elements.extend(create_invoice_header(order_items.iloc[0], styles, doc_type))
    
    # Create and add items table, get calculated totals
    table, subtotal, vat, grand_total = create_invoice_items_table(order_items, styles)
    elements.append(table)
    
    # Add vertical space after table
    elements.append(Spacer(1, 12))
    
    # Add totals section
    elements.append(Paragraph(f"Total excl. VAT: {subtotal:.2f} €", styles["RightBold"]))
    elements.append(Paragraph(f"VAT (21%): {vat:.2f} €", styles["RightBold"]))
    elements.append(Paragraph(f"Total incl. VAT: {grand_total:.2f} €", styles["RightBold"]))
    
    # Add vertical space at the end
    elements.append(Spacer(1, 24))
    
    # Create the PDF document
    document = SimpleDocTemplate(output_path, pagesize=A4, leftMargin=30, rightMargin=30, topMargin=30, bottomMargin=30)
    
    # Build PDF with all elements
    document.build(elements)
    
    return output_path

# =========================
# Document Creation
# =========================

#
# ------------------------------------------------------------
# Generate the internal Word doc used by shop staff for picking
# ------------------------------------------------------------
def create_order_document(order_id, order_items):
    """
    Creates a Word document for an order containing instructions and item details with images.

    Args:
        order_id: ID of the order
        order_items (pandas.DataFrame): DataFrame containing order items
    Returns:
        docx.Document: Generated Word document
    """
    # Create a new Word document
    document = Document()
    
    # Add order title with order ID
    document.add_heading(f"Order {int(float(order_id))}", level=1)
    
    # Resolve which physical shop will fulfil the order so we can address
    # the notification e‑mail correctly.
    shop = order_items.iloc[0]["shop"]
    receiver_email = get_shop_info(shop, 'email')
    
    # Add email header information
    document.add_paragraph(f"From: eshop@winnersport.lt")
    document.add_paragraph(f"To: {receiver_email}")
    
    # Decide what kind of hand‑off: courier/locker vs. in‑store pickup.
    # This determines which instructions the staff will see.
    delivery_method = order_items.iloc[0]["delivery_method"]
    payment_status = order_items.iloc[0]["payment_status"]
    
    # Add instructions heading
    document.add_heading("Instructions:", level=2)
    
    # Add specific instructions based on delivery method and payment status
    if delivery_method in ["by_courier", "parcel_locker"]:
        document.add_paragraph("Please pack the goods securely for delivery. Ensure the correct items and quantities are included.")
    elif delivery_method == "pickup_in_store":
        if payment_status == "received":
            document.add_paragraph("Customer will pick up the order. Provide invoice and ask for a signature.")
        else:
            document.add_paragraph("Customer will pick up the order. A Proforma Invoice is attached.")
    
    # Add items heading
    document.add_heading("Items in Order:", level=2)
    
    # -----------------------------------------------------------------
    # Loop over every product line to embed its description + thumbnail
    # -----------------------------------------------------------------
    for _, item in order_items.iterrows():
        # Extract item details
        item_name = item['item_name']
        size = item.get('size', 'UNI')
        barcode = item.get('barcode', '')
        
        # Format barcode as integer if possible
        if pd.notna(barcode) and str(barcode).strip():
            barcode = f"({int(float(barcode))})"
        else:
            barcode = ''
        
        # Add item description
        document.add_paragraph(f" - Item: {item_name}, size {size}, {barcode} ({item['ordered_quantity']} pcs)")
        
        # Get and process product image
        image_url = item.get("image_url", "")
        image = download_and_process_image(image_url)
        
        # Gracefully degrade: if the image URL is broken, leave a note instead
        if image is None:
            error_msg = f"Could not download image: {image_url}"
            attach_image_to_document(document, None, order_id, error_msg)
        else:
            # Attach image to document
            attach_image_to_document(document, image, order_id)
    
    return document

# =========================
# Order Processing Logic
# =========================

#
# ------------------------------------------------------------
# Produce per-order PDFs: invoices and shipping labels as needed
# ------------------------------------------------------------
def generate_order_pdf_attachments(order_id, order_items, order_folder):
    """
    Generates PDF attachments (invoices, shipping labels) for an order.

    Args:
        order_id: ID of the order
        order_items (pandas.DataFrame): DataFrame containing order items
        order_folder (str): Path to folder where PDF files should be saved
    Returns:
        list: List of paths to generated PDF files
    """
    # List to store paths to generated PDF files
    pdf_files = []
    
    # Convert order_id to integer for filename
    order_num = int(float(order_id))
    
    # Get delivery and payment information
    delivery_method = order_items.iloc[0]["delivery_method"]
    payment_status = order_items.iloc[0]["payment_status"]
    
    # Out‑of‑store deliveries always need an *invoice* + per‑item shipping
    # label(s).  COD info gets baked into the label later if needed.
    if delivery_method in ["by_courier", "parcel_locker"]:
        # Generate invoice PDF
        invoice_pdf = os.path.join(order_folder, f"order_{order_num}_invoice.pdf")
        generate_invoice_pdf(order_num, order_items, invoice_pdf, doc_type="Invoice")
        pdf_files.append(invoice_pdf)
        
        # Generate shipping label for each item
        for _, item in order_items.iterrows():
            item_id = int(item["item_id"])
            
            # Create base filename for label
            label_base = f"order_{order_num}_item_{item_id}_label"
            
            # Add COD (Cash On Delivery) to filename if payment not received
            if payment_status != "received":
                label_base += "_COD"
            
            # Generate label PDF
            label_pdf = os.path.join(order_folder, f"{label_base}.pdf")
            generate_shipping_label_pdf(label_pdf, item, payment_status)
            pdf_files.append(label_pdf)
    
    # Store pickups don’t require courier labels.  We generate either
    # a fiscal invoice (paid) or a pro‑forma (unpaid) instead.
    elif delivery_method == "pickup_in_store":
        if payment_status == "received":
            # Generate regular invoice if payment received
            invoice_pdf = os.path.join(order_folder, f"order_{order_num}_invoice.pdf")
            generate_invoice_pdf(order_num, order_items, invoice_pdf, doc_type="Invoice")
            pdf_files.append(invoice_pdf)
        else:
            # Generate proforma invoice if payment not yet received
            proforma_pdf = os.path.join(order_folder, f"order_{order_num}_proforma_invoice.pdf")
            generate_invoice_pdf(order_num, order_items, proforma_pdf, doc_type="Proforma Invoice")
            pdf_files.append(proforma_pdf)
    
    return pdf_files

#
# ------------------------------------------------------------
# High‑level wrapper that handles one order from start to finish
# ------------------------------------------------------------
def process_individual_order(order_id, order_items, orders_dataframe):
    """
    Processes a single order, generating all necessary documents.

    Args:
        order_id: ID of the order to process
        order_items (pandas.DataFrame): DataFrame containing items for this order
        orders_dataframe (pandas.DataFrame): Complete orders DataFrame (to mark as processed)
    """
    # Mark early to avoid double‑processing even if later steps error out
    # Mark the order as processed in the DataFrame
    mark_order_as_processed(orders_dataframe, order_id)
    
    # Create folder structure for order attachments
    attachments_folder = "order_attachments"
    os.makedirs(attachments_folder, exist_ok=True) # Create main attachments folder if it doesn't exist
    
    # Create folder for this specific order
    order_folder = os.path.join(attachments_folder, f"order_{int(float(order_id))}")
    os.makedirs(order_folder, exist_ok=True) # Create order-specific folder if it doesn't exist
    
    # Word doc doubles as both packing slip and e‑mail body template
    # Create main order document (Word format)
    document = create_order_document(order_id, order_items)
    
    # Generate PDF attachments (invoices, shipping labels, etc.)
    pdf_files = generate_order_pdf_attachments(order_id, order_items, order_folder)
    
    # Add list of attachments to the document if any were created
    if pdf_files:
        document.add_heading("Attachments:", level=2)
        for pdf in pdf_files:
            document.add_paragraph(f"- {os.path.basename(pdf)}")
    
    # Save the Word document
    doc_filename = os.path.join(order_folder, f"Order_{int(float(order_id))}.docx")
    document.save(doc_filename)

# =========================
# Main Entrypoint
# =========================

#
# ------------------------------------------------------------
# Entry point: iterate through Excel and trigger processing
# ------------------------------------------------------------
def process_all_orders():
    """
    Main function that orchestrates the order processing workflow.
    Loads orders from Excel, processes unprocessed orders, and updates the Excel file.
    """
    # NOTE: The Excel sheet acts as the single source of truth.  We load it,
    # mark rows as processed *in‑memory*, and re‑write the same file at the
    # end so that a second run won’t duplicate documents.
    # Define input Excel filename
    filename = "orders.xlsx"
    
    print(f"📄 Loading Excel file: {filename}")
    
    try:
        # Load orders from Excel file
        orders_dataframe = load_orders_from_excel(filename)
        print(f"✅ Loaded {len(orders_dataframe)} rows from Excel")
    except Exception as exception:
        # Handle Excel loading errors
        print(f"❌ Failed to load Excel file: {exception}")
        return
    
    # Check if required columns exist in the Excel file
    if "order_id" not in orders_dataframe.columns or "order_processed" not in orders_dataframe.columns:
        print("❌ Required columns missing in the Excel file.")
        print(f"📊 Columns found: {list(orders_dataframe.columns)}")
        return
    
    # ---------------------------------------------------------------
    # Main loop – executed once per *unique* order_id in the sheet
    # ---------------------------------------------------------------
    # Group orders by order_id and process each order
    for order_id, order_items in orders_dataframe.groupby("order_id"):
        print(f"\n🔍 Checking order {order_id}...")
        
        # Skip already processed orders
        if is_order_already_processed(order_items):
            print(f"⏭️ Skipping already processed order {order_id}")
            continue
        
        print(f"⚙️ Processing order {order_id} with {len(order_items)} items...")
        
        try:
            # Process the order
            process_individual_order(order_id, order_items, orders_dataframe)
            print(f"✅ Order {order_id} processed successfully.")
        except Exception as exception:
            # Handle processing errors
            print(f"❌ Error processing order {order_id}: {exception}")
    
    try:
        # Save updated Excel file with processed orders marked
        orders_dataframe.to_excel(filename, index=False)
        print(f"\n💾 Updated Excel file saved: {filename}")
    except Exception as exception:
        # Handle Excel saving errors
        print(f"❌ Failed to save updated Excel file: {exception}")

# Check if script is run directly (not imported)
if __name__ == "__main__":
    print("✅ Script started...")
    process_all_orders()
